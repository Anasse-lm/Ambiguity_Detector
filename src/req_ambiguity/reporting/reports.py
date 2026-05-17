import io
import json
import sqlite3
from docx import Document
from docx.shared import Pt, Inches

def get_session_data(session_id):
    db_path = "outputs/sessions/session_log.db"
    with sqlite3.connect(db_path) as conn:
        conn.row_factory = sqlite3.Row
        events = [dict(row) for row in conn.execute('SELECT * FROM events WHERE session_id = ? ORDER BY timestamp ASC', (session_id,))]
        stories = [dict(row) for row in conn.execute('SELECT * FROM stories WHERE session_id = ? ORDER BY submitted_at ASC', (session_id,))]
        sessions = [dict(row) for row in conn.execute('SELECT * FROM sessions WHERE session_id = ?', (session_id,))]
    return events, stories, sessions[0] if sessions else None

def per_story_report(session_id, story_id):
    events, stories, _ = get_session_data(session_id)
    story = next((s for s in stories if s['story_id'] == story_id), None)
    if not story:
        return b""
    
    doc = Document()
    doc.add_heading('Refinement Report', 0)
    doc.add_paragraph(f"Session ID: {session_id}")
    doc.add_paragraph(f"Story ID: {story_id}")
    doc.add_paragraph(f"Status: {story['review_status']}")
    
    doc.add_heading('Original Story', level=1)
    doc.add_paragraph(story['original_text'])
    
    pipeline_outputs = json.loads(story['pipeline_outputs_json']) if story['pipeline_outputs_json'] else {}
    
    if 'detection' in pipeline_outputs:
        doc.add_heading('Detection Results', level=1)
        for label, prob in pipeline_outputs['detection'].items():
            doc.add_paragraph(f"{label}: {prob:.4f}")
            
    if 'xai' in pipeline_outputs:
        doc.add_heading('Explanations', level=1)
        for label, exp in pipeline_outputs['xai'].items():
            doc.add_heading(label, level=2)
            doc.add_paragraph("Top evidence tokens:")
            for tok in exp.get('top_evidence_tokens', []):
                doc.add_paragraph(f"{tok[0]} ({tok[1]:.4f})", style='List Bullet')
            
    if 'refinement' in pipeline_outputs:
        doc.add_heading('Refined Story', level=1)
        doc.add_paragraph(pipeline_outputs['refinement'].get('refined_story', ''))
        
        doc.add_heading('Placeholders Inserted', level=2)
        for p in pipeline_outputs['refinement'].get('placeholders_used', []):
            doc.add_paragraph(p, style='List Bullet')
            
        doc.add_heading('Clarification Questions', level=2)
        for q in pipeline_outputs['refinement'].get('clarification_questions', []):
            doc.add_paragraph(q, style='List Number')
            
    if 'verification' in pipeline_outputs:
        doc.add_heading('Verification (Aggregate Delta)', level=1)
        agg_delta = pipeline_outputs['verification'].get('aggregate_delta', 0.0)
        doc.add_paragraph(f"{agg_delta:.4f}")
        
    f = io.BytesIO()
    doc.save(f)
    return f.getvalue()

def session_summary_report(session_id):
    events, stories, session = get_session_data(session_id)
    doc = Document()
    doc.add_heading('Session Summary Report', 0)
    if session:
        doc.add_paragraph(f"Session ID: {session_id}")
        doc.add_paragraph(f"Input Mode: {session['input_mode']}")
        doc.add_paragraph(f"Started At: {session['started_at']}")
        
    doc.add_heading('Stories Overview', level=1)
    doc.add_paragraph(f"Total Stories Processed: {len(stories)}")
    
    accepted = len([s for s in stories if s['review_status'] == 'accepted'])
    skipped = len([s for s in stories if s['review_status'] == 'skipped'])
    regenerated = len([e for e in events if e['event_type'] == 'REFINEMENT_REGENERATED'])
    
    doc.add_paragraph(f"Accepted: {accepted}")
    doc.add_paragraph(f"Skipped: {skipped}")
    doc.add_paragraph(f"Total Regenerations: {regenerated}")
    
    f = io.BytesIO()
    doc.save(f)
    return f.getvalue()

def clarification_questions_report(session_id):
    events, stories, _ = get_session_data(session_id)
    doc = Document()
    doc.add_heading('Clarification Questions', 0)
    
    for story in stories:
        pipeline = json.loads(story['pipeline_outputs_json']) if story['pipeline_outputs_json'] else {}
        ref = pipeline.get('refinement', {})
        questions = ref.get('clarification_questions', [])
        if questions:
            doc.add_heading(f"Story: {story['original_text'][:50]}...", level=1)
            for q in questions:
                doc.add_paragraph(q, style='List Number')
            doc.add_paragraph("Answer: ________________________________________________________\n")
            
    f = io.BytesIO()
    doc.save(f)
    return f.getvalue()

def refined_requirements_report(session_id):
    events, stories, _ = get_session_data(session_id)
    doc = Document()
    doc.add_heading('Refined Requirements Specification', 0)
    doc.add_paragraph("This document contains all accepted stories with ambiguity placeholders preserved.")
    
    for story in stories:
        if story['review_status'] == 'accepted':
            pipeline = json.loads(story['pipeline_outputs_json']) if story['pipeline_outputs_json'] else {}
            ref = pipeline.get('refinement', {})
            refined_text = ref.get('refined_story', '')
            if refined_text:
                doc.add_paragraph(refined_text, style='List Bullet')
                
    f = io.BytesIO()
    doc.save(f)
    return f.getvalue()
